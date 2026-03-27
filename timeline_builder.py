#!/usr/bin/env python3
"""
Litigation Timeline Builder
Parses scanned PDF claim files, OCRs each page, detects document boundaries,
extracts dates and descriptions, and outputs a chronological .docx timeline.

Usage:
    python timeline_builder.py <input.pdf> [--output timeline.docx] [--dpi 200] [--workers 4]
    python timeline_builder.py <input.pdf> --llm-summarize   # Use Claude for better summaries (requires ANTHROPIC_API_KEY)
"""

import argparse
import io
import re
import sys
import json
import os
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

import dateparser
import fitz
import pytesseract
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image


DATE_PATTERNS = [
    re.compile(r'\b(\d{1,2}/\d{1,2}/\d{4})\b'),
    re.compile(r'\b(\d{4}-\d{2}-\d{2})\b'),
    re.compile(r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b', re.IGNORECASE),
    re.compile(r'\b(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b', re.IGNORECASE),
    re.compile(r'\b((?:Mon|Tue|Wed|Thu|Fri|Sat|Sun)\w*,?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b', re.IGNORECASE),
    re.compile(r'\b(\d{1,2}/\d{1,2}/\d{2})\b'),
]

DOCUMENT_START_PATTERNS = [
    re.compile(r'^Property Loss Notice', re.MULTILINE | re.IGNORECASE),
    re.compile(r'^From:\s+.+', re.MULTILINE),
    re.compile(r'VERMONT MUTUAL INSURANCE', re.IGNORECASE),
    re.compile(r'NORTHERN SECURITY INSURANCE', re.IGNORECASE),
    re.compile(r'GRANITE MUTUAL INSURANCE', re.IGNORECASE),
    re.compile(r'^Dear\s+', re.MULTILINE),
    re.compile(r'^RE:\s+', re.MULTILINE | re.IGNORECASE),
    re.compile(r'^SUBJECT:\s+', re.MULTILINE | re.IGNORECASE),
    re.compile(r'CERTIFICATE OF MAILING', re.IGNORECASE),
    re.compile(r'PROOF OF LOSS', re.IGNORECASE),
    re.compile(r'SWORN STATEMENT', re.IGNORECASE),
    re.compile(r'RESERVATION OF RIGHTS', re.IGNORECASE),
    re.compile(r'SCOPE OF LOSS', re.IGNORECASE),
    re.compile(r'INDEPENDENT ADJUSTER', re.IGNORECASE),
    re.compile(r'^INVOICE', re.MULTILINE | re.IGNORECASE),
    re.compile(r'CLAIM\s*(?:NUMBER|#|:)', re.IGNORECASE),
    re.compile(r'SUBROGATION', re.IGNORECASE),
    re.compile(r'EXAMINATION UNDER OATH', re.IGNORECASE),
    re.compile(r'DEMAND LETTER', re.IGNORECASE),
    re.compile(r'COMPLAINT', re.IGNORECASE),
    re.compile(r'SUMMONS', re.IGNORECASE),
    re.compile(r'MOTION\s+(?:TO|FOR)', re.IGNORECASE),
    re.compile(r'(?:GUARD|ARBELLA|AMICA|LIBERTY|TRAVELERS|HANOVER|NORFOLK)\s+INSURANCE', re.IGNORECASE),
    re.compile(r'Independent Claims Service', re.IGNORECASE),
    re.compile(r'ACTIVITY LOG', re.IGNORECASE),
    re.compile(r'PAYMENT SUMMARY', re.IGNORECASE),
    re.compile(r'Page\s*:\s*1\b', re.IGNORECASE),
    re.compile(r'\bPage\s+1\s+of\s+\d+', re.IGNORECASE),
]

CONTINUATION_PATTERNS = [
    re.compile(r'^CONTINUED\s*[-—]', re.MULTILINE | re.IGNORECASE),
    re.compile(r'Page\s*:\s*(?:[2-9]|\d{2,})\b'),
    re.compile(r'\bPage\s+(?:[2-9]|\d{2,})\s+of\s+\d+', re.IGNORECASE),
    re.compile(r'^DESCRIPTION\s+QUANTITY\s+UNIT', re.MULTILINE | re.IGNORECASE),
]

ESTIMATE_HEADER = re.compile(r'Independent Claims Service', re.IGNORECASE)
ESTIMATE_CONTINUATION = re.compile(r'CONTINUED\s*[-—]', re.IGNORECASE)


def ocr_page(args):
    """OCR a single page. Designed to run in a subprocess."""
    pdf_path, page_num, dpi = args
    doc = fitz.open(pdf_path)
    page = doc[page_num]
    pix = page.get_pixmap(dpi=dpi)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    text = pytesseract.image_to_string(img)
    doc.close()
    return page_num, text


def ocr_all_pages(pdf_path, dpi=200, workers=4, cache_path=None):
    """OCR every page of the PDF, with optional caching."""
    if cache_path and os.path.exists(cache_path):
        print(f"Loading cached OCR from {cache_path}")
        with open(cache_path, "r") as f:
            return json.load(f)

    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    doc.close()

    print(f"OCR processing {total_pages} pages at {dpi} DPI with {workers} workers...")
    pages_text = [""] * total_pages
    args_list = [(pdf_path, i, dpi) for i in range(total_pages)]

    completed = 0
    with ProcessPoolExecutor(max_workers=workers) as executor:
        futures = {executor.submit(ocr_page, args): args[1] for args in args_list}
        for future in as_completed(futures):
            page_num, text = future.result()
            pages_text[page_num] = text
            completed += 1
            if completed % 50 == 0 or completed == total_pages:
                pct = (completed / total_pages) * 100
                print(f"  [{completed}/{total_pages}] {pct:.1f}% complete")

    if cache_path:
        print(f"Caching OCR results to {cache_path}")
        with open(cache_path, "w") as f:
            json.dump(pages_text, f)

    return pages_text


def is_estimate_page(text):
    """Check if page is part of the ICS damage estimate report."""
    return bool(ESTIMATE_HEADER.search(text))


def classify_page(text, page_num):
    """Classify a page and detect if it's the start of a new document."""
    text_stripped = text.strip()
    if not text_stripped:
        return "blank", False

    is_continuation = any(p.search(text_stripped) for p in CONTINUATION_PATTERNS)
    is_new_doc = any(p.search(text_stripped) for p in DOCUMENT_START_PATTERNS)

    if is_continuation and not is_new_doc:
        return "continuation", False

    if is_new_doc and not is_continuation:
        return "new_document", True

    if is_new_doc and is_continuation:
        first_lines = "\n".join(text_stripped.split("\n")[:5])
        if any(p.search(first_lines) for p in DOCUMENT_START_PATTERNS):
            return "new_document", True
        return "continuation", False

    return "content", False


def extract_dates(text):
    """Extract all dates from text, return as sorted list of datetime objects."""
    dates = []
    for pattern in DATE_PATTERNS:
        for match in pattern.finditer(text):
            date_str = match.group(1)
            parsed = dateparser.parse(
                date_str,
                settings={
                    "PREFER_DAY_OF_MONTH": "first",
                    "STRICT_PARSING": True,
                    "REQUIRE_PARTS": ["year"],
                }
            )
            if parsed and 2020 <= parsed.year <= 2030:
                dates.append(parsed)
    return sorted(set(dates))


def identify_document_type(text):
    """Attempt to identify the type of document from its content."""
    first_500 = text[:500].lower()

    if "property loss notice" in first_500:
        return "Property Loss Notice"
    if "proof of loss" in first_500:
        return "Proof of Loss"
    if "reservation of rights" in first_500:
        return "Reservation of Rights Letter"
    if "sworn statement" in first_500:
        return "Sworn Statement"
    if "examination under oath" in first_500:
        return "Examination Under Oath"
    if "subrogation" in first_500:
        return "Subrogation Document"
    if "scope of loss" in first_500:
        return "Scope of Loss Report"
    if "activity log" in first_500:
        return "Activity Log"
    if "payment summary" in first_500:
        return "Payment Summary"
    if "invoice" in first_500:
        return "Invoice"
    if "independent claims service" in first_500:
        return "Damage Estimate / Inspection Report"
    if re.search(r"from:\s+.+\nse?nt:", first_500):
        return "Email Correspondence"
    if "dear " in first_500:
        return "Letter"
    if "complaint" in first_500:
        return "Complaint"
    if "summons" in first_500:
        return "Summons"
    if re.search(r"motion\s+(?:to|for)", first_500):
        return "Motion"
    if "demand" in first_500:
        return "Demand Letter"
    if re.search(r"vermont mutual|northern security|granite mutual", first_500):
        return "Insurance Company Correspondence"
    if "guard insurance" in first_500:
        return "Agent Correspondence"
    if "certificate" in first_500:
        return "Certificate"
    if re.search(r"\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}", first_500):
        return "Adjuster Notes / Log Entry"
    if "photo" in first_500 or "date taken" in first_500:
        return "Photo Documentation"

    return "Document"


# Boilerplate patterns to skip when extracting summaries
BOILERPLATE = (
    r"confidential|privileged|intended solely for|if you received this.*error|"
    r"page\s+\d+\s+of\s+\d+|please notify the sender|delete the original|"
    r"do not forward|dissemination.?distribution|prohibited|"
    r"vermont mutual|northern security|granite mutual|since 1828|"
    r"^from:\s|^to:\s|^sent:\s|^subject:\s|^cc:\s|^reply.?to:|"
    r"^\s*\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}\s*-"  # timestamp log prefix
)
FORM_LABEL_RE = re.compile(r"\s*\[\s*[\w\s,]+\]\s*")  # form labels like [number], [Manager, Assignment] (throughout string)
BOILERPLATE_RE = re.compile(BOILERPLATE, re.IGNORECASE | re.MULTILINE)

# Email header lines (stop parsing body when we hit these again)
EMAIL_HEADER_KEYS = ("from:", "to:", "sent:", "subject:", "cc:", "bcc:", "reply-to:", "attachments:")


def _first_substantive_content(text, max_chars=150):
    """Extract first non-boilerplate, substantive sentence or phrase."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:25]:
        if len(line) < 20:
            continue
        if BOILERPLATE_RE.search(line):
            continue
        # Skip lines that are mostly punctuation or numbers
        if re.match(r"^[\d\s\.\-\$\,]+$", line):
            continue
        # Skip form field labels (OCR artifacts)
        if re.match(r"^\[\s*[\w\s,]+\]\s*$", line):
            continue
        # Strip trailing [label] artifacts
        cleaned = FORM_LABEL_RE.sub("", line).strip()
        if len(cleaned) >= 20:
            return cleaned[:max_chars]
    return None


def _extract_claim_numbers(text):
    """Extract claim/policy numbers from text."""
    claims = []
    for m in re.finditer(r"(?:Claim|File)\s*(?:#|:)?\s*([A-Z0-9\-]+)", text[:800], re.IGNORECASE):
        val = m.group(1).strip()
        if len(val) >= 4 and val not in claims:
            claims.append(val)
    for m in re.finditer(r"Policy:\s*([A-Z0-9\-]+)", text[:500], re.IGNORECASE):
        val = m.group(1).strip()
        if len(val) >= 4 and val not in claims:
            claims.append(val)
    return claims[:2]  # max 2


def _extract_amounts(text):
    """Extract notable dollar amounts."""
    amounts = []
    for m in re.finditer(r"\$[\d,]+(?:\.\d{2})?", text[:1500]):
        amounts.append(m.group(0))
    return amounts[:2]


def generate_description(text, doc_type):
    """Generate a brief, informative description of the document."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    claim_nums = _extract_claim_numbers(text)
    claim_str = f" [{', '.join(claim_nums)}]" if claim_nums else ""

    if doc_type == "Email Correspondence":
        from_line = ""
        to_line = ""
        subject_line = ""
        body_start = None

        for i, line in enumerate(lines):
            lower = line.lower()
            if lower.startswith("from:"):
                from_line = line[5:].strip()
                name_match = re.match(r"(.+?)\s*<", from_line)
                if name_match:
                    from_line = name_match.group(1).strip()
            elif lower.startswith("subject:"):
                subject_line = line[8:].strip()
            elif lower.startswith("to:"):
                to_line = line[3:].strip()[:50]
            elif any(lower.startswith(k) for k in EMAIL_HEADER_KEYS):
                continue
            elif len(line) > 25 and not BOILERPLATE_RE.search(line):
                body_start = line
                break

        parts = []
        if from_line:
            parts.append(f"From {from_line}")
        if subject_line:
            parts.append(f"re: {subject_line[:80]}")
        if body_start and not subject_line:
            parts.append(body_start[:80])
        if not parts:
            return f"Email{claim_str}" if claim_str else "Email"

        result = "; ".join(parts)
        return f"{result}{claim_str}" if claim_str else result

    if doc_type == "Adjuster Notes / Log Entry":
        for line in lines[:8]:
            if re.match(r"\d{1,2}/\d{1,2}/\d{4}", line):
                content = re.sub(r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}[^\-]*-\s*", "", line).strip()
                if len(content) > 20:
                    return f"{content[:140]}{claim_str}" if claim_str else content[:140]
        return f"Adjuster log entry{claim_str}" if claim_str else "Adjuster log entry"

    if doc_type in ("Damage Estimate / Inspection Report", "Photo Documentation"):
        for line in lines[:15]:
            if "date taken" in line.lower():
                return f"{line[:130]}{claim_str}" if claim_str else line[:130]
            if re.search(r"basement|kitchen|living room|bedroom|roof|exterior", line, re.I):
                return f"Scope: {line[:100]}{claim_str}" if claim_str else f"Scope: {line[:100]}"
        amounts = _extract_amounts(text)
        amt_str = f" (RCV: {amounts[0]})" if amounts else ""
        return f"{doc_type}{amt_str}{claim_str}" if (amt_str or claim_str) else doc_type

    if doc_type == "Property Loss Notice":
        date_match = re.search(r"Date of Loss:\s*(\S+)", text[:600])
        agent_match = re.search(r"Agent[,\s]+Policy[,\s]+Caller\s*\n\s*(\S.+?)(?:\n|$)", text[:400], re.DOTALL)
        agent = agent_match.group(1).strip()[:40] if agent_match else None

        parts = [f"Claim {claim_nums[0]}" if claim_nums else "Claim filed"]
        if date_match:
            parts.append(f"DOL {date_match.group(1)}")
        if agent:
            parts.append(f"via {agent}")
        return " - ".join(parts)

    if doc_type == "Insurance Company Correspondence":
        for line in lines:
            if line.lower().startswith("re:"):
                return f"RE: {line[3:].strip()[:120]}{claim_str}" if claim_str else f"RE: {line[3:].strip()[:120]}"
        content = _first_substantive_content(text, 120)
        return f"{content}{claim_str}" if content and claim_str else (content or f"Insurance letter{claim_str}")

    if doc_type == "Letter":
        for line in lines:
            if line.lower().startswith("re:"):
                return f"RE: {line[3:].strip()[:120]}{claim_str}" if claim_str else f"RE: {line[3:].strip()[:120]}"
        content = _first_substantive_content(text, 120)
        return f"{content}{claim_str}" if content and claim_str else (content or f"Letter{claim_str}")

    if doc_type == "Invoice":
        for m in re.finditer(r"(?:invoice|inv)\s*#?\s*(\S+)", text[:300], re.I):
            return f"Invoice {m.group(1)}" + (f" - {_extract_amounts(text)[0]}" if _extract_amounts(text) else "")
        amounts = _extract_amounts(text)
        return f"Invoice{claim_str} ({amounts[0]})" if amounts and claim_str else (f"Invoice ({amounts[0]})" if amounts else f"Invoice{claim_str}")

    if doc_type == "Payment Summary":
        amounts = _extract_amounts(text)
        content = _first_substantive_content(text, 100)
        if amounts:
            return f"Payment summary: {amounts[0]}" + (f" - {content}" if content else "") + claim_str
        return f"Payment summary{claim_str}" if claim_str else (content or "Payment summary")

    if doc_type == "Activity Log":
        content = _first_substantive_content(text, 120)
        return f"{content}{claim_str}" if content and claim_str else (f"Activity log{claim_str}" if claim_str else "Activity log")

    if doc_type == "Document":
        content = _first_substantive_content(text, 130)
        if content:
            return f"{content}{claim_str}" if claim_str else content
        for line in lines[:10]:
            if 30 < len(line) < 200 and not BOILERPLATE_RE.search(line):
                cleaned = FORM_LABEL_RE.sub("", line).strip()
                if len(cleaned) >= 25:
                    return f"{cleaned[:120]}{claim_str}" if claim_str else cleaned[:120]

    content = _first_substantive_content(text, 120)
    return f"{content}{claim_str}" if content and claim_str else (content or doc_type)


def llm_summarize_document(text, doc_type, rule_description, max_chars=4000):
    """
    Use Claude to generate a concise 1-2 sentence summary. Optional enhancement.
    Requires anthropic package and ANTHROPIC_API_KEY env var.
    """
    if not HAS_ANTHROPIC:
        return None
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return None

    excerpt = text[:max_chars] if len(text) > max_chars else text
    if not excerpt.strip():
        return None

    prompt = f"""You are summarizing litigation/insurance claim documents for a timeline. Given this document excerpt and its type, write a single concise summary (1-2 sentences, max ~25 words) that captures the key content: what it is, who it's from/to, and the main point.

Document type: {doc_type}
Current rule-based description: {rule_description}

Document excerpt:
{excerpt}

Summary:"""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=150,
            messages=[{"role": "user", "content": prompt}],
        )
        summary = resp.content[0].text.strip()
        return summary[:200] if summary else None
    except Exception:
        return None


def segment_documents(pages_text):
    """
    Segment the OCR'd pages into distinct documents.
    Returns a list of document dicts with page ranges, text, type, dates, description.
    """
    documents = []
    current_doc = None
    in_estimate_block = False

    for i, text in enumerate(pages_text):
        text_stripped = text.strip()

        if not text_stripped:
            if current_doc:
                current_doc["pages"].append(i)
                current_doc["text"] += "\n\n[blank page]\n\n"
            continue

        is_estimate = is_estimate_page(text)

        if is_estimate:
            if ESTIMATE_CONTINUATION.search(text) and in_estimate_block and current_doc:
                current_doc["pages"].append(i)
                current_doc["text"] += "\n\n" + text
                continue
            elif not in_estimate_block or not current_doc:
                if current_doc:
                    documents.append(current_doc)
                in_estimate_block = True
                current_doc = {
                    "pages": [i],
                    "text": text,
                    "start_page": i + 1,
                }
                continue
            else:
                page_match = re.search(r'Page\s*:\s*(\d+)', text)
                if page_match and int(page_match.group(1)) == 1:
                    if current_doc:
                        documents.append(current_doc)
                    current_doc = {
                        "pages": [i],
                        "text": text,
                        "start_page": i + 1,
                    }
                    continue
                current_doc["pages"].append(i)
                current_doc["text"] += "\n\n" + text
                continue

        if in_estimate_block:
            in_estimate_block = False

        page_type, is_new = classify_page(text, i)

        if is_new or current_doc is None:
            if current_doc:
                documents.append(current_doc)
            current_doc = {
                "pages": [i],
                "text": text,
                "start_page": i + 1,
            }
        else:
            if current_doc:
                current_doc["pages"].append(i)
                current_doc["text"] += "\n\n" + text
            else:
                current_doc = {
                    "pages": [i],
                    "text": text,
                    "start_page": i + 1,
                }

    if current_doc:
        documents.append(current_doc)

    for doc in documents:
        doc["end_page"] = max(doc["pages"]) + 1
        doc["page_count"] = len(doc["pages"])
        doc["type"] = identify_document_type(doc["text"])
        doc["dates"] = extract_dates(doc["text"])
        doc["primary_date"] = doc["dates"][0] if doc["dates"] else None
        raw_desc = generate_description(doc["text"], doc["type"])
        doc["description"] = FORM_LABEL_RE.sub(" ", raw_desc).replace("  ", " ").strip()

    return documents


def merge_estimate_blocks(documents):
    """Merge consecutive ICS estimate pages into single documents when appropriate."""
    merged = []
    i = 0
    while i < len(documents):
        doc = documents[i]
        if doc["type"] == "Damage Estimate / Inspection Report":
            while (i + 1 < len(documents)
                   and documents[i + 1]["type"] == "Damage Estimate / Inspection Report"
                   and documents[i + 1]["start_page"] == doc["end_page"] + 1):
                next_doc = documents[i + 1]
                doc["pages"].extend(next_doc["pages"])
                doc["text"] += "\n\n" + next_doc["text"]
                doc["end_page"] = next_doc["end_page"]
                doc["page_count"] = len(doc["pages"])
                all_dates = doc["dates"] + next_doc["dates"]
                doc["dates"] = sorted(set(all_dates))
                doc["primary_date"] = doc["dates"][0] if doc["dates"] else None
                i += 1
            doc["description"] = f"Damage Estimate / Inspection Report ({doc['page_count']} pages, pp. {doc['start_page']}-{doc['end_page']})"
        merged.append(doc)
        i += 1
    return merged


def build_docx(documents, output_path, pdf_name):
    """Build the .docx timeline from segmented documents."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    title = doc.add_heading("Litigation Timeline", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Source: {pdf_name}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = meta2.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    meta3 = doc.add_paragraph()
    meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta3.add_run(f"Documents identified: {len(documents)}")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    dated_docs = [d for d in documents if d["primary_date"]]
    undated_docs = [d for d in documents if not d["primary_date"]]

    dated_docs.sort(key=lambda d: d["primary_date"])

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    headers = ["#", "Date", "Document Type", "Description", "Pages"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for idx, d in enumerate(dated_docs, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = d["primary_date"].strftime("%m/%d/%Y")
        row_cells[2].text = d["type"]
        row_cells[3].text = d["description"]
        if d["start_page"] == d["end_page"]:
            row_cells[4].text = str(d["start_page"])
        else:
            row_cells[4].text = f"{d['start_page']}-{d['end_page']}"

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    if undated_docs:
        doc.add_paragraph()
        doc.add_heading("Undated Documents", level=1)
        undated_table = doc.add_table(rows=1, cols=4)
        undated_table.style = "Light Grid Accent 1"

        header_cells = undated_table.rows[0].cells
        headers = ["#", "Document Type", "Description", "Pages"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for idx, d in enumerate(undated_docs, 1):
            row_cells = undated_table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = d["type"]
            row_cells[2].text = d["description"]
            if d["start_page"] == d["end_page"]:
                row_cells[3].text = str(d["start_page"])
            else:
                row_cells[3].text = f"{d['start_page']}-{d['end_page']}"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    widths = [Inches(0.4), Inches(1.0), Inches(1.8), Inches(3.0), Inches(0.8)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.save(output_path)
    print(f"\nTimeline saved to: {output_path}")
    print(f"  Dated documents: {len(dated_docs)}")
    print(f"  Undated documents: {len(undated_docs)}")
    print(f"  Total: {len(documents)}")


def main():
    parser = argparse.ArgumentParser(description="Build a litigation timeline from a scanned PDF claim file.")
    parser.add_argument("pdf", help="Path to the input PDF file")
    parser.add_argument("--output", "-o", help="Output .docx path (default: <input>_timeline.docx)")
    parser.add_argument("--dpi", type=int, default=200, help="OCR resolution in DPI (default: 200)")
    parser.add_argument("--workers", "-w", type=int, default=4, help="Number of parallel OCR workers (default: 4)")
    parser.add_argument("--cache", action="store_true", help="Cache OCR results for re-runs")
    parser.add_argument("--summary", action="store_true", help="Print document summary to stdout")
    parser.add_argument("--llm-summarize", action="store_true", help="Use Claude API to enhance descriptions (requires ANTHROPIC_API_KEY and anthropic package)")
    args = parser.parse_args()

    pdf_path = Path(args.pdf).resolve()
    if not pdf_path.exists():
        print(f"Error: {pdf_path} not found")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output).resolve()
    else:
        output_path = pdf_path.with_name(pdf_path.stem + "_timeline.docx")

    cache_path = None
    if args.cache:
        cache_path = str(pdf_path.with_suffix(".ocr_cache.json"))

    print(f"Input:   {pdf_path}")
    print(f"Output:  {output_path}")
    print(f"DPI:     {args.dpi}")
    print(f"Workers: {args.workers}")
    print()

    pages_text = ocr_all_pages(str(pdf_path), dpi=args.dpi, workers=args.workers, cache_path=cache_path)

    print(f"\nSegmenting {len(pages_text)} pages into documents...")
    documents = segment_documents(pages_text)
    print(f"Found {len(documents)} raw document segments")

    documents = merge_estimate_blocks(documents)
    print(f"After merging estimate blocks: {len(documents)} documents")

    if args.llm_summarize:
        if not HAS_ANTHROPIC:
            print("Error: --llm-summarize requires the anthropic package. Run: pip install anthropic")
            sys.exit(1)
        if not os.environ.get("ANTHROPIC_API_KEY"):
            print("Error: --llm-summarize requires ANTHROPIC_API_KEY environment variable")
            sys.exit(1)
        print(f"\nEnhancing descriptions with Claude ({len(documents)} documents)...")
        for i, d in enumerate(documents):
            enhanced = llm_summarize_document(d["text"], d["type"], d["description"])
            if enhanced:
                d["description"] = enhanced
            if (i + 1) % 25 == 0 or i == len(documents) - 1:
                print(f"  [{i + 1}/{len(documents)}] summarized")
            time.sleep(0.3)  # gentle rate limiting

    if args.summary:
        print("\n" + "=" * 80)
        print("DOCUMENT SUMMARY")
        print("=" * 80)
        for i, d in enumerate(documents, 1):
            date_str = d["primary_date"].strftime("%m/%d/%Y") if d["primary_date"] else "No date"
            print(f"\n{i}. [{date_str}] {d['type']}")
            print(f"   Pages: {d['start_page']}-{d['end_page']} ({d['page_count']} pages)")
            print(f"   {d['description']}")

    build_docx(documents, str(output_path), pdf_path.name)


if __name__ == "__main__":
    main()
